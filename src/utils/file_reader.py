"""
文件读取模块

支持 Word (.docx)、PDF (.pdf)、Markdown (.md) 文件的文本提取
"""

import os
from pathlib import Path
from typing import Optional, List, Dict, Any

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

from loguru import logger


class FileReader:
    """文件读取器"""

    # 支持的文件格式
    SUPPORTED_FORMATS = {
        ".docx": "word",
        ".doc": "word",
        ".pdf": "pdf",
        ".md": "markdown",
        ".txt": "text",
    }

    @staticmethod
    def read_file(file_path: str) -> Optional[str]:
        """
        读取文件内容

        Args:
            file_path (str): 文件路径

        Returns:
            Optional[str]: 文件内容，如果失败返回None
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None

        suffix = file_path.suffix.lower()

        if suffix not in FileReader.SUPPORTED_FORMATS:
            logger.error(f"不支持的文件格式: {suffix}")
            return None

        file_type = FileReader.SUPPORTED_FORMATS[suffix]

        try:
            if file_type == "word":
                return FileReader._read_docx(file_path)
            elif file_type == "pdf":
                return FileReader._read_pdf(file_path)
            elif file_type in ["markdown", "text"]:
                return FileReader._read_text(file_path)
        except Exception as e:
            logger.error(f"读取文件失败 {file_path}: {str(e)}")
            return None

    @staticmethod
    def _read_docx(file_path: Path) -> Optional[str]:
        """
        读取Word文档内容
        """
        if Document is None:
            logger.error("python-docx not installed")
            return None

        try:
            doc = Document(file_path)
            text_content = []

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))

            return "\n".join(text_content)
        except Exception as e:
            logger.error(f"读取DOCX失败: {str(e)}")
            return None

    @staticmethod
    def _read_pdf(file_path: Path) -> Optional[str]:
        """
        读取PDF文件内容
        """
        if PdfReader is None:
            logger.error("PyPDF2 not installed")
            return None

        try:
            pdf_reader = PdfReader(file_path)
            text_content = []

            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    text_content.append(text)

            return "\n".join(text_content)
        except Exception as e:
            logger.error(f"读取PDF失败: {str(e)}")
            return None

    @staticmethod
    def _read_text(file_path: Path) -> Optional[str]:
        """
        读取文本文件
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"读取文本文件失败: {str(e)}")
                return None

    @staticmethod
    def read_directory(
        dir_path: str,
        file_types: Optional[List[str]] = None,
        recursive: bool = True,
    ) -> Dict[str, str]:
        """
        批量读取目录中的文件

        Args:
            dir_path (str): 目录路径
            file_types (Optional[List[str]]): 要读取的文件类型，如[".md", ".docx"]
            recursive (bool): 是否递归读取子目录

        Returns:
            Dict[str, str]: {文件路径: 文件内容}
        """
        dir_path = Path(dir_path)

        if not dir_path.exists():
            logger.error(f"目录不存在: {dir_path}")
            return {}

        if file_types is None:
            file_types = list(FileReader.SUPPORTED_FORMATS.keys())

        file_contents = {}
        pattern = "**/*" if recursive else "*"

        for file_path in dir_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in file_types:
                content = FileReader.read_file(str(file_path))
                if content:
                    file_contents[str(file_path)] = content
                    logger.info(f"读取文件: {file_path}")

        logger.info(f"共读取 {len(file_contents)} 个文件")
        return file_contents
